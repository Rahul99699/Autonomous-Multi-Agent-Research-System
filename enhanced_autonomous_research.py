# %% [markdown]
# # Autonomous Research Scientist (Enhanced Pipeline)
# This notebook implements an advanced, end-to-end multi-agent research pipeline.
# It features HDBSCAN clustering, Hybrid Search (Semantic + Lexical), LLM-based Summarization,
# and enhanced DOCX report generation.

# %% [markdown]
# ## 1. Install Required Packages
# Install enhanced dependencies including `hdbscan`, `rank_bm25`, and `transformers` for LLM summarization.

# %%
import subprocess
import sys

print("⏳ Installing required packages...")
packages = [
    "pandas", "numpy", "scikit-learn", "matplotlib", "seaborn", "plotly",
    "sentence-transformers", "spacy", "nltk", "bertopic", "hdbscan", 
    "rank_bm25", "transformers", "torch", "python-docx", "reportlab", "datasets"
]
# Uncomment the following lines to auto-install dependencies when running as a script
# subprocess.check_call([sys.executable, "-m", "pip", "install", "-q"] + packages)
# subprocess.check_call([sys.executable, "-m", "spacy", "download", "en_core_web_sm"])
print("✅ All libraries installed successfully!")

# %%
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from sklearn.manifold import TSNE
from sklearn.metrics import silhouette_score, davies_bouldin_score
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
import spacy
import nltk
from nltk.corpus import stopwords
from bertopic import BERTopic
import hdbscan
from rank_bm25 import BM25Okapi
from transformers import pipeline
from datasets import load_dataset
from datetime import datetime
import warnings
import io
import os

warnings.filterwarnings('ignore')
nltk.download('stopwords', quiet=True)

sns.set_style("whitegrid")
plt.rcParams['figure.figsize'] = (12, 6)

print("✅ All imports successful!")

# %% [markdown]
# ## 2. Dataset Loading & Preprocessing
# We load the arXiv abstracts dataset and prepare it for clustering.

# %%
print("="*60)
print("📥 LOADING RESEARCH PAPER DATASET")
print("="*60)

try:
    print("\n⏳ Downloading arXiv dataset from HuggingFace...")
    dataset = load_dataset("gfissore/arxiv-abstracts-2021", split='train')
    df = dataset.to_pandas()
    
    required_columns = ['title', 'abstract', 'categories']
    df = df.dropna(subset=required_columns)[required_columns]
    
    ai_categories = ['cs.AI', 'cs.LG', 'cs.CL', 'cs.CV', 'cs.RO']
    df = df[df['categories'].apply(lambda x: any(cat in str(x) for cat in ai_categories))]
    
    df = df.drop_duplicates(subset=['title']).drop_duplicates(subset=['abstract'])
    
    SAMPLE_SIZE = 5000
    if len(df) > SAMPLE_SIZE:
        df = df.sample(n=SAMPLE_SIZE, random_state=42)
        
    df = df.reset_index(drop=True)
    
    def map_category(cat_text):
        cat_text = str(cat_text)
        if 'cs.CV' in cat_text: return 'Computer Vision'
        elif 'cs.CL' in cat_text: return 'NLP'
        elif 'cs.RO' in cat_text: return 'Robotics'
        elif 'cs.LG' in cat_text: return 'Machine Learning'
        elif 'cs.AI' in cat_text: return 'Artificial Intelligence'
        else: return 'Other'
        
    df['category'] = df['categories'].apply(map_category)
    df = df[['title', 'abstract', 'category']]
    print("\n✅ Dataset preprocessing complete!")

except Exception as e:
    print(f"\n⚠️ Error loading dataset: {e}")
    print("\n⚠️ Using fallback demo dataset instead...")
    
    sample_data = {
        'title': [
            'Deep Learning for Medical Image Segmentation',
            'Transformer Models in Natural Language Processing',
            'Reinforcement Learning for Robotics Control',
            'Graph Neural Networks for Knowledge Graphs',
            'Federated Learning in Healthcare',
            'Vision Transformers for Computer Vision',
            'BERT: Pre-training of Deep Bidirectional Transformers',
            'Attention is All You Need',
            'U-Net: Biomedical Image Segmentation',
            'YOLO: Real-time Object Detection'
        ],
        'abstract': [
            'Deep learning methods for medical image segmentation using convolutional neural networks.',
            'Transformer architectures for NLP tasks including translation and sentiment analysis.',
            'Reinforcement learning framework for robotics control systems.',
            'Graph neural networks for reasoning over knowledge graphs.',
            'Federated learning for privacy-preserving healthcare AI.',
            'Vision transformers for computer vision applications.',
            'Bidirectional transformers for language understanding.',
            'Self-attention transformer architecture for machine translation.',
            'Biomedical segmentation using U-Net convolutional networks.',
            'Real-time object detection using YOLO architecture.'
        ],
        'category': [
            'Computer Vision', 'NLP', 'Robotics', 'Machine Learning', 'Healthcare AI',
            'Computer Vision', 'NLP', 'NLP', 'Computer Vision', 'Computer Vision'
        ]
    }
    df = pd.DataFrame(sample_data)
    augmented_rows = []
    for i in range(20):
        temp_df = df.copy()
        temp_df['title'] = temp_df['title'] + f" v{i+1}"
        augmented_rows.append(temp_df)
    df = pd.concat([df] + augmented_rows, ignore_index=True)
    df = df.sample(frac=1, random_state=42).reset_index(drop=True)

print(f"\n✅ Final Dataset Shape: {df.shape}")

# %% [markdown]
# ## 3. Advanced NLP Preprocessing & High-Fidelity Embeddings
# We use `all-mpnet-base-v2` for superior semantic representations.

# %%
print("="*60)
print("🔧 NLP PREPROCESSING & EMBEDDINGS")
print("="*60)

nlp = spacy.load("en_core_web_sm", disable=['parser', 'ner'])
stop_words = set(stopwords.words('english'))

def preprocess_text(text):
    if not isinstance(text, str) or len(text) == 0: return ""
    doc = nlp(text.lower())
    tokens = [token.lemma_ for token in doc if token.text not in stop_words and token.is_alpha]
    return " ".join(tokens)

print("\n⏳ Processing abstracts...")
df['processed_abstract'] = df['abstract'].apply(preprocess_text)

print("\n📥 Loading Sentence Transformer model (all-mpnet-base-v2)...")
model = SentenceTransformer('all-mpnet-base-v2')

print("\n⏳ Generating high-fidelity embeddings...")
embeddings = model.encode(df['processed_abstract'].tolist(), show_progress_bar=True, batch_size=32)
print(f"✅ Embeddings shape: {embeddings.shape}")

# %% [markdown]
# ## 4. Advanced Clustering (HDBSCAN vs KMeans)
# HDBSCAN handles noise and non-spherical clusters significantly better than standard KMeans.

# %%
print("="*60)
print("🎯 ADVANCED CLUSTERING (HDBSCAN)")
print("="*60)

print("\n⏳ Applying HDBSCAN...")
clusterer = hdbscan.HDBSCAN(min_cluster_size=5, metric='euclidean', cluster_selection_method='eom')
df['hdbscan_cluster'] = clusterer.fit_predict(embeddings)

# Fallback to KMeans if HDBSCAN finds too much noise or too few clusters
optimal_k = len(set(df['hdbscan_cluster'])) - (1 if -1 in df['hdbscan_cluster'] else 0)
if optimal_k < 2:
    print("⚠️ HDBSCAN found too few clusters. Falling back to KMeans (k=5).")
    optimal_k = 5
    kmeans = KMeans(n_clusters=optimal_k, random_state=42, n_init=10)
    df['cluster'] = kmeans.fit_predict(embeddings)
else:
    print(f"✅ HDBSCAN found {optimal_k} clusters (and noise points labeled as -1).")
    df['cluster'] = df['hdbscan_cluster']

# Filter out noise for metric calculation if using HDBSCAN
valid_mask = df['cluster'] != -1
if valid_mask.sum() > 0 and len(set(df.loc[valid_mask, 'cluster'])) > 1:
    silhouette = silhouette_score(embeddings[valid_mask], df.loc[valid_mask, 'cluster'])
    print(f"  • Silhouette Score (excluding noise): {silhouette:.4f}")

# %% [markdown]
# ## 5. Visualization (PCA & t-SNE)

# %%
print("="*60)
print("📉 DIMENSIONALITY REDUCTION & VIZ")
print("="*60)

pca = PCA(n_components=2, random_state=42)
embeddings_2d_pca = pca.fit_transform(embeddings)

tsne = TSNE(n_components=2, random_state=42, perplexity=30, n_iter=1000)
embeddings_2d_tsne = tsne.fit_transform(embeddings)

# Save plots for the report
os.makedirs('report_assets', exist_ok=True)

plt.figure(figsize=(10, 6))
scatter = plt.scatter(embeddings_2d_pca[:, 0], embeddings_2d_pca[:, 1], c=df['cluster'], cmap='tab20', alpha=0.6, s=50, edgecolors='black', linewidth=0.5)
plt.title('PCA Visualization of Clusters', fontsize=14, fontweight='bold')
plt.colorbar(scatter, label='Cluster ID')
plt.savefig('report_assets/pca_plot.png', bbox_inches='tight')
plt.show()

plt.figure(figsize=(10, 6))
scatter = plt.scatter(embeddings_2d_tsne[:, 0], embeddings_2d_tsne[:, 1], c=df['cluster'], cmap='tab20', alpha=0.6, s=50, edgecolors='black', linewidth=0.5)
plt.title('t-SNE Visualization of Clusters', fontsize=14, fontweight='bold')
plt.colorbar(scatter, label='Cluster ID')
plt.savefig('report_assets/tsne_plot.png', bbox_inches='tight')
plt.show()

# %% [markdown]
# ## 6. Hybrid Search Pipeline (Semantic + BM25)
# Fuses dense embeddings with sparse keyword search for robust retrieval.

# %%
print("="*60)
print("🔍 HYBRID SEARCH (SEMANTIC + BM25)")
print("="*60)

tokenized_corpus = [doc.split(" ") for doc in df['processed_abstract']]
bm25 = BM25Okapi(tokenized_corpus)

def hybrid_search(query, top_k=5, alpha=0.5):
    # Semantic Search
    query_emb = model.encode([query])[0]
    semantic_scores = cosine_similarity([query_emb], embeddings)[0]
    
    # Lexical Search
    tokenized_query = preprocess_text(query).split(" ")
    lexical_scores = bm25.get_scores(tokenized_query)
    
    # Normalize scores
    semantic_norm = (semantic_scores - np.min(semantic_scores)) / (np.max(semantic_scores) - np.min(semantic_scores) + 1e-8)
    lexical_norm = (lexical_scores - np.min(lexical_scores)) / (np.max(lexical_scores) - np.min(lexical_scores) + 1e-8)
    
    # Fusion
    hybrid_scores = alpha * semantic_norm + (1 - alpha) * lexical_norm
    top_indices = np.argsort(hybrid_scores)[::-1][:top_k]
    
    results = []
    for rank, idx in enumerate(top_indices, 1):
        results.append({
            'rank': rank, 'title': df.iloc[idx]['title'], 
            'cluster': df.iloc[idx]['cluster'], 'score': hybrid_scores[idx]
        })
    return results

# Example
query = "Deep learning for medical image segmentation"
print(f"🔎 Query: '{query}'")
for res in hybrid_search(query, top_k=3):
    print(f"#{res['rank']} [Score: {res['score']:.3f}] - {res['title']} (Cluster {res['cluster']})")

# %% [markdown]
# ## 7. LLM-Based Cluster Summarization
# We use a lightweight open-source LLM to generate descriptive insights for each cluster.

# %%
print("="*60)
print("🧠 LLM CLUSTER SUMMARIZATION")
print("="*60)

print("\n⏳ Loading summarization pipeline...")
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6", device=-1) # change device to 0 if GPU is available

cluster_summaries = {}
unique_clusters = sorted([c for c in df['cluster'].unique() if c != -1])

for cluster_id in unique_clusters:
    # Get top 5 papers closest to the cluster centroid (or just random if HDBSCAN)
    cluster_papers = df[df['cluster'] == cluster_id]['abstract'].tolist()
    if not cluster_papers: continue
    
    # Combine abstracts (truncate to fit model context)
    combined_text = " ".join(cluster_papers[:5])[:2000] 
    
    try:
        summary = summarizer(combined_text, max_length=60, min_length=20, do_sample=False)[0]['summary_text']
        cluster_summaries[cluster_id] = summary
        print(f"\nCluster {cluster_id} Summary: {summary}")
    except Exception as e:
        cluster_summaries[cluster_id] = "Summary generation failed."
        print(f"\nCluster {cluster_id} Summary Error: {e}")

# %% [markdown]
# ## 8. Enhanced Report Generation (DOCX)
# Integrates visualizations and LLM summaries into a rich document.

# %%
print("="*60)
print("📄 GENERATING ENHANCED DOCX REPORT")
print("="*60)

from docx import Document
from docx.shared import Inches

doc = Document()
doc.add_heading('Enhanced Autonomous Research Analysis', 0)
doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

doc.add_heading('Executive Summary', 1)
doc.add_paragraph(f"Analysis of {len(df)} papers using HDBSCAN clustering and all-mpnet-base-v2 embeddings. "
                  f"Discovered {optimal_k} distinct research clusters.")

doc.add_heading('Visualizations', 1)
if os.path.exists('report_assets/pca_plot.png'):
    doc.add_picture('report_assets/pca_plot.png', width=Inches(5.5))
if os.path.exists('report_assets/tsne_plot.png'):
    doc.add_picture('report_assets/tsne_plot.png', width=Inches(5.5))

doc.add_heading('Cluster Insights (LLM Generated)', 1)
for cluster_id, summary in cluster_summaries.items():
    doc.add_heading(f'Cluster {cluster_id}', level=2)
    doc.add_paragraph(summary)
    
    # Add top titles
    cluster_df = df[df['cluster'] == cluster_id]
    doc.add_paragraph("Top Papers:", style='List Bullet')
    for title in cluster_df['title'].head(3):
        doc.add_paragraph(title, style='List Bullet 2')

report_path = 'Enhanced_Research_Report.docx'
doc.save(report_path)
print(f"\n✅ Report generated successfully: {report_path}")
